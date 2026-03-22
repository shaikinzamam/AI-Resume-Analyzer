import re
import json
import base64
import os
import subprocess
import tempfile
import requests
from io import BytesIO

OLLAMA_URL = "http://localhost:11434/api/generate"
PREFERRED_MODELS = [
    "phi3:mini",
    "mistral:7b-instruct-q4_0",
    "mistral:7b-instruct-q4_K_M",
    "mistral",
]

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1: Parse the ORIGINAL resume to extract real structural data
# ─────────────────────────────────────────────────────────────────────────────

def parse_original_resume(text: str) -> dict:
    """
    Extract real name, contact, companies, dates, education from resume text.
    This data is ALWAYS used — even if Ollama fails.
    Returns a base dict that Ollama output is merged into.
    """
    lines = [l.strip() for l in text.replace("\r\n", "\n").split("\n") if l.strip()]

    result = {
        "name": "",
        "contact": "",
        "email": "",
        "phone": "",
        "linkedin": "",
        "github": "",
        "location": "",
        "tagline": "",
        "summary": "",
        "skills": [],
        "experience": [],
        "projects": [],
        "education": [],
        "certifications": [],
    }

    # ── Name: usually the first non-empty line, no special chars
    for line in lines[:5]:
        if (not re.search(r'[@|:/\d]', line)
                and len(line.split()) <= 5
                and len(line) > 3):
            result["name"] = line.title() if line.isupper() else line
            break

    # ── Contact details
    for line in lines[:15]:
        email_m = re.search(r'[\w.+-]+@[\w.-]+\.\w+', line)
        phone_m = re.search(r'[\+\d][\d\s\-().]{8,}', line)
        linkedin_m = re.search(r'linkedin\.com/in/[\w-]+', line, re.I)
        github_m = re.search(r'github\.com/[\w-]+', line, re.I)

        if email_m and not result["email"]:
            result["email"] = email_m.group()
        if phone_m and not result["phone"]:
            result["phone"] = phone_m.group().strip()
        if linkedin_m and not result["linkedin"]:
            result["linkedin"] = "https://" + linkedin_m.group()
        if github_m and not result["github"]:
            result["github"] = "https://" + github_m.group()

    # ── Location (look for city/state patterns)
    for line in lines[:20]:
        if re.search(r'\b(andhra|telangana|hyderabad|bangalore|chennai|mumbai|delhi|IN|India)\b', line, re.I):
            loc_m = re.search(r'([A-Za-z\s]+,\s*[A-Z]{2,})', line)
            if loc_m:
                result["location"] = loc_m.group()

    # ── Skills section
    skill_section = False
    for i, line in enumerate(lines):
        up = line.upper()
        if re.match(r'(SKILLS?|TECHNICAL SKILLS?|CORE SKILLS?|KEY SKILLS?)', up):
            skill_section = True
            continue
        if skill_section:
            # Stop at next section header
            if re.match(r'(EXPERIENCE|EDUCATION|PROJECT|CERTIF|SUMMARY|WORK)', up):
                break
            # Parse comma/bullet separated skills
            clean = re.sub(r'^[-•*·:]\s*', '', line)
            for sk in re.split(r'[,|•·/]', clean):
                sk = sk.strip()
                if 2 < len(sk) < 40 and not re.search(r'^\d+$', sk):
                    result["skills"].append(sk)

    # ── Experience blocks
    exp_section = False
    current_job = None
    for line in lines:
        up = line.upper()
        if re.match(r'(EXPERIENCE|WORK HISTORY|EMPLOYMENT)', up):
            exp_section = True
            continue
        if exp_section and re.match(r'(PROJECTS?|EDUCATION|CERTIF|ACHIEVEMENTS?)', up):
            if current_job:
                result["experience"].append(current_job)
                current_job = None
            exp_section = False
            continue
        if not exp_section:
            continue

        is_bullet = bool(re.match(r'^[-•*·]', line))
        has_date  = bool(re.search(r'\b(20\d\d|present|current|jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\b', line, re.I))
        has_pipe  = "|" in line

        if (has_date or has_pipe) and not is_bullet and len(line) > 5:
            if current_job:
                result["experience"].append(current_job)
            # Parse: "Title | Company | Dates" or "Title at Company (dates)"
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 3:
                title, company, duration = parts[0], parts[1], parts[2]
            elif len(parts) == 2:
                title, company = parts[0], parts[1]
                # Extract date from company part if present
                date_m = re.search(r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[\w\s,–-]+(?:20\d\d|present))', company, re.I)
                if date_m:
                    duration = date_m.group()
                    company  = company.replace(duration, "").strip(" |–-")
                else:
                    duration = ""
            else:
                title, company, duration = line, "", ""
            current_job = {"title": title, "company": company, "duration": duration, "bullets": []}
        elif is_bullet and current_job is not None:
            bullet = re.sub(r'^[-•*·]\s*', '', line)
            current_job["bullets"].append(bullet)

    if current_job:
        result["experience"].append(current_job)

    # ── Projects
    proj_section = False
    current_proj = None
    for line in lines:
        up = line.upper()
        if re.match(r'PROJECTS?', up):
            proj_section = True
            continue
        if proj_section and re.match(r'(EDUCATION|CERTIF|ACHIEVEMENTS?|SKILLS?)', up):
            if current_proj:
                result["projects"].append(current_proj)
                current_proj = None
            proj_section = False
            continue
        if not proj_section:
            continue

        is_bullet = bool(re.match(r'^[-•*·]', line))
        is_link   = bool(re.search(r'github\.com|https?://', line, re.I))

        if not is_bullet and not is_link and len(line.split()) >= 2:
            if current_proj:
                result["projects"].append(current_proj)
            parts = [p.strip() for p in line.split("|")]
            name = parts[0]
            tech = parts[1] if len(parts) > 1 else ""
            # Also handle [tech stack] in brackets
            bracket_m = re.search(r'\[([^\]]+)\]', name)
            if bracket_m:
                tech = bracket_m.group(1)
                name = name.replace(bracket_m.group(), "").strip()
            current_proj = {"name": name, "tech": tech, "bullets": [], "link": ""}
        elif is_bullet and current_proj is not None:
            bullet = re.sub(r'^[-•*·]\s*', '', line)
            current_proj["bullets"].append(bullet)
        elif is_link and current_proj is not None:
            link_m = re.search(r'https?://\S+', line)
            if link_m:
                current_proj["link"] = link_m.group()

    if current_proj:
        result["projects"].append(current_proj)

    # ── Education
    edu_section = False
    for line in lines:
        up = line.upper()
        if re.match(r'EDUCATION', up):
            edu_section = True
            continue
        if edu_section and re.match(r'(CERTIF|ACHIEVEMENTS?|SKILLS?|PROJECTS?)', up):
            edu_section = False
            continue
        if edu_section and len(line) > 5:
            clean = re.sub(r'^[-•*·]\s*', '', line)
            result["education"].append(clean)

    # ── Certifications
    cert_section = False
    for line in lines:
        up = line.upper()
        if re.match(r'CERTIF', up):
            cert_section = True
            continue
        if cert_section and re.match(r'(EDUCATION|SKILLS?|PROJECTS?|EXPERIENCE)', up):
            cert_section = False
            continue
        if cert_section and len(line) > 3:
            clean = re.sub(r'^[-•*·]\s*', '', line)
            result["certifications"].append(clean)

    # ── Build contact line
    parts = []
    if result["email"]:    parts.append(result["email"])
    if result["phone"]:    parts.append(result["phone"])
    if result["linkedin"]: parts.append(result["linkedin"])
    if result["github"]:   parts.append(result["github"])
    if result["location"]: parts.append(result["location"])
    result["contact"] = "   |   ".join(parts)

    return result


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2: Locally improve bullets when Ollama is unavailable
# ─────────────────────────────────────────────────────────────────────────────

WEAK_VERBS = {
    r'^worked on\b':          'Developed',
    r'^used\b':               'Leveraged',
    r'^helped\b':             'Supported',
    r'^was responsible for\b':'Owned',
    r'^responsible for\b':    'Owned',
    r'^involved in\b':        'Contributed to',
    r'^participated in\b':    'Contributed to',
    r'^assisted\b':           'Supported',
    r'^did\b':                'Executed',
    r'^made\b':               'Built',
    r'^tried\b':              'Implemented',
}

def _improve_bullet_locally(bullet: str, jd_skills: list) -> str:
    """Replace weak verbs; fix tech capitalisation."""
    improved = bullet.strip()
    for pattern, replacement in WEAK_VERBS.items():
        if re.match(pattern, improved, re.I):
            improved = re.sub(pattern, replacement, improved, flags=re.I, count=1)
            improved = improved[0].upper() + improved[1:]
            break
    # Fix common tech capitalizations
    fixes = {
        r'\bpython\b':'Python', r'\bfastapi\b':'FastAPI', r'\bflask\b':'Flask',
        r'\bdjango\b':'Django', r'\bdocker\b':'Docker', r'\baws\b':'AWS',
        r'\bsql\b':'SQL', r'\bnosql\b':'NoSQL', r'\bmongodb\b':'MongoDB',
        r'\bpostgresql\b':'PostgreSQL', r'\bgit\b':'Git', r'\bgithub\b':'GitHub',
        r'\blangchain\b':'LangChain', r'\bopenai\b':'OpenAI', r'\bllm\b':'LLM',
        r'\brag\b':'RAG', r'\bfaiss\b':'FAISS', r'\bollama\b':'Ollama',
        r'\bapi\b':'API', r'\brest api\b':'REST API', r'\bn8n\b':'n8n',
        r'\bjavascript\b':'JavaScript', r'\btypescript\b':'TypeScript',
        r'\breact\b':'React', r'\bnumpy\b':'NumPy', r'\bpandas\b':'Pandas',
    }
    for pattern, proper in fixes.items():
        improved = re.sub(pattern, proper, improved, flags=re.I)
    return improved


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3: Ollama — plain text rewrite (fast, no JSON)
# ─────────────────────────────────────────────────────────────────────────────

def _get_model() -> str:
    try:
        resp = requests.get("http://localhost:11434/api/tags", timeout=5)
        installed = {m["name"] for m in resp.json().get("models", [])}
        for m in PREFERRED_MODELS:
            if m in installed:
                return m
    except Exception:
        pass
    return PREFERRED_MODELS[-1]


REWRITE_PROMPT = """\
You are a senior ATS resume writer. Rewrite ONLY the bullet points and professional summary below.

STRICT RULES:
1. Keep every real company name, job title, date, and project name EXACTLY as given.
2. Every bullet MUST start with a strong action verb: Built, Developed, Architected, Designed,
   Implemented, Deployed, Automated, Optimised, Integrated, Delivered, Led, Reduced, Increased.
3. Add a specific metric to EVERY bullet: %, count, time saved, users, requests/sec, etc.
   If the original has no metric, invent a plausible one (e.g. "reducing latency by ~30%").
4. Inject keywords from the JD naturally into bullets and summary.
5. Write a 4-sentence professional summary tailored to the JD.
6. Output ONLY in this exact format — no JSON, no extra text:

TAGLINE
[one-line title e.g. "AI / Python Developer | LLM Applications | RAG Systems"]

SUMMARY
[4 sentences]

EXPERIENCE
[Job Title] | [Company] | [Dates]
- [rewritten bullet with metric]
- [rewritten bullet with metric]

PROJECTS
[Project Name] | [Tech]
- [rewritten bullet with metric]

---
ORIGINAL RESUME:
{resume}

JOB DESCRIPTION:
{jd}

Rewrite now:"""


def _stream_ollama(prompt: str) -> str:
    model = _get_model()
    parts = []
    with requests.post(
        OLLAMA_URL,
        json={
            "model": model,
            "prompt": prompt,
            "stream": True,
            "options": {"temperature": 0.3, "num_predict": 1800, "num_ctx": 4096},
        },
        stream=True,
        timeout=(10, 480),
    ) as resp:
        resp.raise_for_status()
        for line in resp.iter_lines():
            if not line:
                continue
            try:
                chunk = json.loads(line)
                parts.append(chunk.get("response", ""))
                if chunk.get("done"):
                    break
            except Exception:
                continue
    return "".join(parts)


def _parse_ollama_output(text: str) -> dict:
    """Parse Ollama plain-text output into structured dict."""
    result = {"tagline": "", "summary": "", "experience": [], "projects": []}
    current_section = None
    current_job = None
    current_proj = None
    summary_lines = []

    for line in text.replace("\r\n", "\n").split("\n"):
        s = line.strip()
        up = s.upper()

        if up == "TAGLINE":
            current_section = "tagline"; continue
        if up == "SUMMARY":
            current_section = "summary"; continue
        if up == "EXPERIENCE":
            if summary_lines:
                result["summary"] = " ".join(summary_lines)
                summary_lines = []
            current_section = "experience"; continue
        if up == "PROJECTS":
            if current_job:
                result["experience"].append(current_job)
                current_job = None
            current_section = "projects"; continue
        if s in ("---", ""):
            continue

        if current_section == "tagline" and not result["tagline"]:
            result["tagline"] = s

        elif current_section == "summary":
            if s:
                summary_lines.append(s)

        elif current_section == "experience":
            is_bullet = s.startswith("-")
            has_date  = bool(re.search(r'20\d\d|present', s, re.I))
            has_pipe  = "|" in s
            if (has_date or has_pipe) and not is_bullet:
                if current_job:
                    result["experience"].append(current_job)
                parts = [p.strip() for p in s.split("|")]
                title    = parts[0] if len(parts) > 0 else s
                company  = parts[1] if len(parts) > 1 else ""
                duration = parts[2] if len(parts) > 2 else ""
                current_job = {"title": title, "company": company, "duration": duration, "bullets": []}
            elif is_bullet and current_job:
                current_job["bullets"].append(s[1:].strip())

        elif current_section == "projects":
            is_bullet = s.startswith("-")
            if not is_bullet and s and not s.startswith("---"):
                if current_proj:
                    result["projects"].append(current_proj)
                parts = [p.strip() for p in s.split("|")]
                current_proj = {"name": parts[0], "tech": parts[1] if len(parts) > 1 else "", "bullets": []}
            elif is_bullet and current_proj:
                current_proj["bullets"].append(s[1:].strip())

    if current_job:
        result["experience"].append(current_job)
    if current_proj:
        result["projects"].append(current_proj)
    if summary_lines and not result["summary"]:
        result["summary"] = " ".join(summary_lines)

    return result


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4: Merge parsed original + Ollama output
# ─────────────────────────────────────────────────────────────────────────────

def _merge(parsed: dict, ollama: dict, jd: str) -> dict:
    """
    Merge strategy:
    - Name, contact, education, certifications → ALWAYS from parsed (real data)
    - Tagline, summary → from Ollama if available, else generate locally
    - Experience bullets → from Ollama if company names match, else locally improved
    - Projects → Ollama bullets if available, else locally improved originals
    """
    merged = dict(parsed)

    # ── Tagline
    if ollama.get("tagline"):
        merged["tagline"] = ollama["tagline"]
    elif merged["skills"]:
        top = merged["skills"][:3]
        merged["tagline"] = " | ".join(top)

    # ── Summary
    if ollama.get("summary") and len(ollama["summary"]) > 50:
        merged["summary"] = ollama["summary"]

    # ── Experience: use Ollama bullets if we got them, but keep real company/title/date
    jd_skills = _extract_jd_skills(jd)
    if ollama.get("experience"):
        ollama_exp = ollama["experience"]
        for i, job in enumerate(merged["experience"]):
            if i < len(ollama_exp) and ollama_exp[i].get("bullets"):
                job["bullets"] = ollama_exp[i]["bullets"]
            else:
                # Improve locally
                job["bullets"] = [_improve_bullet_locally(b, jd_skills) for b in job["bullets"]]
    else:
        for job in merged["experience"]:
            job["bullets"] = [_improve_bullet_locally(b, jd_skills) for b in job["bullets"]]

    # ── Projects: same logic
    if ollama.get("projects"):
        ollama_projs = ollama["projects"]
        for i, proj in enumerate(merged["projects"]):
            if i < len(ollama_projs) and ollama_projs[i].get("bullets"):
                proj["bullets"] = ollama_projs[i]["bullets"]
            else:
                proj["bullets"] = [_improve_bullet_locally(b, jd_skills) for b in proj["bullets"]]
    else:
        for proj in merged["projects"]:
            proj["bullets"] = [_improve_bullet_locally(b, jd_skills) for b in proj["bullets"]]

    return merged


def _extract_jd_skills(jd: str) -> list:
    keywords = ["python","fastapi","django","flask","docker","kubernetes","aws","azure","gcp",
                 "sql","mongodb","postgresql","git","langchain","rag","llm","genai","openai",
                 "machine learning","deep learning","nlp","react","typescript","javascript",
                 "n8n","automation","faiss","ollama","rest api","microservices","ci/cd"]
    jd_lower = jd.lower()
    return [k for k in keywords if k in jd_lower]


# ─────────────────────────────────────────────────────────────────────────────
# STEP 5: Build DOCX via Node.js docx library
# ─────────────────────────────────────────────────────────────────────────────

DOCX_SCRIPT = r"""
const {{
  Document, Packer, Paragraph, TextRun, AlignmentType,
  BorderStyle, LevelFormat, ExternalHyperlink, TabStopType,
  UnderlineType
}} = require("{docx_path}");
const fs = require("fs");

const NAVY   = "1A3A5C";
const DARK   = "111111";
const GREY   = "333333";
const LGREY  = "999999";
const ACCENT = "2563EB";
const PAGE_W = 12240;
const PAGE_H = 15840;
const MARGIN = 864;
const CWIDTH = PAGE_W - MARGIN * 2;

const data = {data_json};

function hr() {{
  return new Paragraph({{
    border: {{ bottom: {{ style: BorderStyle.SINGLE, size: 5, color: "CCCCCC", space: 1 }} }},
    spacing: {{ before: 0, after: 80 }},
    children: [],
  }});
}}

function sectionHead(text) {{
  return [
    new Paragraph({{
      spacing: {{ before: 180, after: 20 }},
      children: [new TextRun({{ text: text.toUpperCase(), bold: true, size: 21, color: NAVY, font: "Calibri" }})],
    }}),
    hr(),
  ];
}}

function nameBlock() {{
  const rows = [
    new Paragraph({{
      alignment: AlignmentType.CENTER,
      spacing: {{ before: 0, after: 30 }},
      children: [new TextRun({{ text: data.name || "Candidate", bold: true, size: 56, font: "Calibri", color: DARK }})],
    }}),
  ];
  if (data.tagline) {{
    rows.push(new Paragraph({{
      alignment: AlignmentType.CENTER,
      spacing: {{ before: 0, after: 40 }},
      children: [new TextRun({{ text: data.tagline, size: 20, font: "Calibri", color: NAVY, italics: true }})],
    }}));
  }}
  // Contact line with clickable links
  const contactRuns = [];
  const addPart = (label, url) => {{
    if (!label) return;
    if (contactRuns.length > 0) {{
      contactRuns.push(new TextRun({{ text: "   |   ", size: 18, font: "Calibri", color: LGREY }}));
    }}
    if (url) {{
      contactRuns.push(new ExternalHyperlink({{ link: url, children: [
        new TextRun({{ text: label, size: 18, font: "Calibri", color: ACCENT,
          underline: {{ type: UnderlineType.SINGLE }} }}),
      ]}}));
    }} else {{
      contactRuns.push(new TextRun({{ text: label, size: 18, font: "Calibri", color: GREY }}));
    }}
  }};
  addPart(data.email, null);
  addPart(data.phone, null);
  if (data.linkedin) addPart("LinkedIn", data.linkedin);
  if (data.github)   addPart("GitHub",   data.github);
  if (data.location) addPart(data.location, null);
  if (contactRuns.length) {{
    rows.push(new Paragraph({{
      alignment: AlignmentType.CENTER,
      spacing: {{ before: 0, after: 60 }},
      children: contactRuns,
    }}));
  }}
  rows.push(hr());
  return rows;
}}

function summaryBlock() {{
  if (!data.summary) return [];
  return [
    ...sectionHead("Professional Summary"),
    new Paragraph({{
      spacing: {{ before: 40, after: 80 }},
      children: [new TextRun({{ text: data.summary, size: 20, font: "Calibri", color: GREY }})],
    }}),
  ];
}}

function skillsBlock() {{
  if (!data.skills || !data.skills.length) return [];
  const half = Math.ceil(data.skills.length / 2);
  const row1 = data.skills.slice(0, half).join("   •   ");
  const row2 = data.skills.slice(half).join("   •   ");
  const paras = [
    new Paragraph({{
      spacing: {{ before: 40, after: 20 }},
      children: [new TextRun({{ text: row1, size: 20, font: "Calibri", color: GREY }})],
    }}),
  ];
  if (row2) paras.push(new Paragraph({{
    spacing: {{ before: 0, after: 80 }},
    children: [new TextRun({{ text: row2, size: 20, font: "Calibri", color: GREY }})],
  }}));
  return [...sectionHead("Core Skills"), ...paras];
}}

function experienceBlock() {{
  if (!data.experience || !data.experience.length) return [];
  const items = [];
  for (const job of data.experience) {{
    const titleParts = [];
    if (job.title)   titleParts.push(new TextRun({{ text: job.title,   bold: true, size: 21, font: "Calibri", color: DARK }}));
    if (job.company) titleParts.push(
      new TextRun({{ text: "  |  ", size: 21, font: "Calibri", color: LGREY }}),
      new TextRun({{ text: job.company, size: 21, font: "Calibri", color: NAVY }}),
    );
    if (job.duration) titleParts.push(
      new TextRun({{ text: "\t" + job.duration, size: 19, font: "Calibri", color: LGREY, italics: true }}),
    );
    items.push(new Paragraph({{
      tabStops: [{{ type: TabStopType.RIGHT, position: CWIDTH }}],
      spacing: {{ before: 120, after: 30 }},
      children: titleParts,
    }}));
    for (const b of (job.bullets || [])) {{
      items.push(new Paragraph({{
        numbering: {{ reference: "bullets", level: 0 }},
        spacing: {{ before: 20, after: 20 }},
        children: [new TextRun({{ text: b, size: 20, font: "Calibri", color: GREY }})],
      }}));
    }}
  }}
  return [...sectionHead("Professional Experience"), ...items];
}}

function projectsBlock() {{
  if (!data.projects || !data.projects.length) return [];
  const items = [];
  for (const proj of data.projects) {{
    const hdr = [];
    hdr.push(new TextRun({{ text: proj.name, bold: true, size: 21, font: "Calibri", color: DARK }}));
    if (proj.tech) {{
      hdr.push(new TextRun({{ text: "  |  ", size: 20, font: "Calibri", color: LGREY }}));
      hdr.push(new TextRun({{ text: proj.tech, size: 19, font: "Calibri", color: GREY, italics: true }}));
    }}
    items.push(new Paragraph({{ spacing: {{ before: 100, after: 20 }}, children: hdr }}));
    for (const b of (proj.bullets || [])) {{
      items.push(new Paragraph({{
        numbering: {{ reference: "bullets", level: 0 }},
        spacing: {{ before: 20, after: 20 }},
        children: [new TextRun({{ text: b, size: 20, font: "Calibri", color: GREY }})],
      }}));
    }}
    if (proj.link) {{
      items.push(new Paragraph({{
        spacing: {{ before: 10, after: 30 }},
        children: [
          new TextRun({{ text: "GitHub: ", size: 18, font: "Calibri", color: LGREY }}),
          new ExternalHyperlink({{ link: proj.link, children: [
            new TextRun({{ text: proj.link, size: 18, font: "Calibri", color: ACCENT,
              underline: {{ type: UnderlineType.SINGLE }} }}),
          ]}}),
        ],
      }}));
    }}
  }}
  return [...sectionHead("Projects"), ...items];
}}

function educationBlock() {{
  if (!data.education || !data.education.length) return [];
  const items = data.education.map(e => new Paragraph({{
    spacing: {{ before: 60, after: 30 }},
    children: [new TextRun({{ text: e, size: 20, font: "Calibri", color: GREY }})],
  }}));
  return [...sectionHead("Education"), ...items];
}}

function certsBlock() {{
  if (!data.certifications || !data.certifications.length) return [];
  const items = data.certifications.map(c => new Paragraph({{
    numbering: {{ reference: "bullets", level: 0 }},
    spacing: {{ before: 20, after: 20 }},
    children: [new TextRun({{ text: c, size: 20, font: "Calibri", color: GREY }})],
  }}));
  return [...sectionHead("Certifications"), ...items];
}}

const doc = new Document({{
  numbering: {{
    config: [{{
      reference: "bullets",
      levels: [{{
        level: 0,
        format: LevelFormat.BULLET,
        text: "•",
        alignment: AlignmentType.LEFT,
        style: {{ paragraph: {{ indent: {{ left: 360, hanging: 220 }} }} }},
      }}],
    }}],
  }},
  styles: {{
    default: {{ document: {{ run: {{ font: "Calibri", size: 20, color: GREY }} }} }},
  }},
  sections: [{{
    properties: {{
      page: {{
        size: {{ width: PAGE_W, height: PAGE_H }},
        margin: {{ top: MARGIN, bottom: MARGIN, left: MARGIN, right: MARGIN }},
      }},
    }},
    children: [
      ...nameBlock(),
      ...summaryBlock(),
      ...skillsBlock(),
      ...experienceBlock(),
      ...projectsBlock(),
      ...educationBlock(),
      ...certsBlock(),
    ],
  }}],
}});

Packer.toBuffer(doc).then(buf => {{
  fs.writeFileSync("{output_path}", buf);
  console.log("OK");
}}).catch(e => {{ console.error(e.message); process.exit(1); }});
"""


def _build_docx_via_node(data: dict) -> bytes:
    """Render the DOCX using Node.js docx library for best formatting."""
    # Try to find docx module via npm
    candidate_paths = [
        "/home/claude/.npm-global/lib/node_modules/docx/dist/index.cjs",
        "/usr/local/lib/node_modules/docx/dist/index.cjs",
        "/usr/lib/node_modules/docx/dist/index.cjs",
    ]
    # Also try `npm root -g`
    docx_path = None
    for p in candidate_paths:
        if os.path.exists(p):
            docx_path = p
            break
    if not docx_path:
        try:
            npm_root = subprocess.check_output(["npm", "root", "-g"], text=True, timeout=5).strip()
            p = os.path.join(npm_root, "docx", "dist", "index.cjs")
            if os.path.exists(p):
                docx_path = p
        except Exception:
            pass
    if not docx_path:
        return _build_docx_python(data)

    with tempfile.TemporaryDirectory() as tmpdir:
        output_path = os.path.join(tmpdir, "resume.docx")
        script_path = os.path.join(tmpdir, "build.cjs")

        script = DOCX_SCRIPT.format(
            docx_path=docx_path,
            data_json=json.dumps(data, ensure_ascii=False),
            output_path=output_path,
        )
        with open(script_path, "w", encoding="utf-8") as f:
            f.write(script)

        result = subprocess.run(
            ["node", script_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0 or not os.path.exists(output_path):
            raise RuntimeError(f"Node build failed: {result.stderr}")

        with open(output_path, "rb") as f:
            return f.read()


def _build_docx_python(data: dict) -> bytes:
    """Python-docx fallback (used only if Node.js unavailable)."""
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDoc()
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.65)
        sec.left_margin = sec.right_margin = Inches(0.75)

    NAVY_C = RGBColor(0x1A, 0x3A, 0x5C)
    GREY_C = RGBColor(0x33, 0x33, 0x33)
    LGREY_C = RGBColor(0xAA, 0xAA, 0xAA)

    def heading(txt):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(txt.upper())
        r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = NAVY_C
        hr = doc.add_paragraph()
        hr.paragraph_format.space_before = Pt(0)
        hr.paragraph_format.space_after = Pt(5)
        rr = hr.add_run("─" * 90); rr.font.size = Pt(6.5); rr.font.color.rgb = LGREY_C

    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(data.get("name") or "Resume")
    r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(0x11,0x11,0x11)
    if data.get("tagline"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.add_run(data["tagline"]).font.size = Pt(10)
    if data.get("contact"):
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(data["contact"]).font.size = Pt(9)

    if data.get("summary"):
        heading("Professional Summary")
        doc.add_paragraph(data["summary"])

    if data.get("skills"):
        heading("Core Skills")
        half = max(1, len(data["skills"]) // 2)
        doc.add_paragraph("   •   ".join(data["skills"][:half]))
        if data["skills"][half:]:
            doc.add_paragraph("   •   ".join(data["skills"][half:]))

    if data.get("experience"):
        heading("Professional Experience")
        for job in data["experience"]:
            p = doc.add_paragraph()
            p.add_run(f"{job.get('title','')}  |  {job.get('company','')}  |  {job.get('duration','')}").bold = True
            for b in job.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(b)

    if data.get("projects"):
        heading("Projects")
        for proj in data["projects"]:
            p = doc.add_paragraph()
            p.add_run(proj.get("name","")).bold = True
            if proj.get("tech"):
                p.add_run(f"  |  {proj['tech']}")
            for b in proj.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(b)
            if proj.get("link"):
                doc.add_paragraph(f"GitHub: {proj['link']}")

    if data.get("education"):
        heading("Education")
        for e in data["education"]:
            doc.add_paragraph(e)

    if data.get("certifications"):
        heading("Certifications")
        for c in data["certifications"]:
            doc.add_paragraph(f"• {c}")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _docx_to_pdf(docx_bytes: bytes):
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            path = os.path.join(tmpdir, "resume.docx")
            with open(path, "wb") as f: f.write(docx_bytes)
            r = subprocess.run(
                ["libreoffice","--headless","--convert-to","pdf","--outdir",tmpdir,path],
                capture_output=True, timeout=60)
            pdf = os.path.join(tmpdir, "resume.pdf")
            if r.returncode == 0 and os.path.exists(pdf):
                with open(pdf, "rb") as pf:
                    return pf.read()
    except Exception:
        pass
    return None


# ─────────────────────────────────────────────────────────────────────────────
# PUBLIC API
# ─────────────────────────────────────────────────────────────────────────────

def build_resume_for_jd(resume_text: str, job_description: str) -> dict:
    """
    Main entry point called by app.py → /api/build-resume
    Returns: { rewritten, docx_b64, pdf_b64, warning, raw_output }
    """
    warning = ""

    # 1. Always parse the real resume first
    parsed = parse_original_resume(resume_text)

    # 2. Try Ollama for better bullets + summary
    raw_output = ""
    ollama_data = {}
    try:
        safe_resume = resume_text[:3000].replace("{", "(").replace("}", ")")
        safe_jd = job_description[:1500].replace("{", "(").replace("}", ")")
        prompt = REWRITE_PROMPT.format(
            resume=safe_resume,
            jd=safe_jd,
        )
        raw_output = _stream_ollama(prompt)
        if raw_output.strip():
            ollama_data = _parse_ollama_output(raw_output)
    except requests.ConnectionError:
        warning = "Ollama offline — used original resume with improved formatting."
    except requests.Timeout:
        warning = "Ollama timed out — used original resume. Try: ollama pull phi3:mini"
    except Exception as e:
        warning = f"Ollama error ({e}) — used original resume."

    # 3. Merge real data + Ollama improvements
    final_data = _merge(parsed, ollama_data, job_description)

    # 4. Build DOCX
    try:
        docx_bytes = _build_docx_via_node(final_data)
        docx_b64 = base64.b64encode(docx_bytes).decode()
    except Exception as e:
        try:
            docx_bytes = _build_docx_python(final_data)
            docx_b64 = base64.b64encode(docx_bytes).decode()
            warning += f" | Node.js unavailable, used python-docx."
        except Exception as e2:
            docx_bytes = b""
            docx_b64 = ""
            warning += f" | DOCX build failed: {e2}"

    # 5. Try PDF
    pdf_b64 = ""
    if docx_bytes:
        pdf_bytes = _docx_to_pdf(docx_bytes)
        if pdf_bytes:
            pdf_b64 = base64.b64encode(pdf_bytes).decode()

    return {
        "rewritten":  final_data,
        "docx_b64":   docx_b64,
        "pdf_b64":    pdf_b64,
        "warning":    warning,
        "raw_output": raw_output,
    }